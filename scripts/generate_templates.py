"""Generate DOCX templates with placeholders for Garant lawyer agent."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def _add_centered_heading(
    doc: Document, text: str, size: int = 14, bold: bool = True
) -> None:
    """Add centered bold paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_placeholder_paragraph(doc: Document, placeholder: str) -> None:
    """Add placeholder line."""
    p = doc.add_paragraph()
    run = p.add_run(placeholder)
    run.italic = True
    run.font.color.rgb = None  # Keep default color
    p.paragraph_format.space_after = Pt(6)


def _add_section_heading(doc: Document, text: str) -> None:
    """Add section heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def create_pretension_template(output_path: Path) -> Path:
    """Create pretension DOCX template."""
    doc = Document()

    _add_placeholder_paragraph(doc, "{наименование_адресат}")
    _add_placeholder_paragraph(doc, "{наименование_адресат_доп}")
    _add_placeholder_paragraph(doc, "От {ФИО_истец}")
    _add_placeholder_paragraph(doc, "Адрес: {адрес_истец}")
    _add_placeholder_paragraph(doc, "Тел.: {телефон_истец}, email: {email_истец}")

    doc.add_paragraph()
    _add_centered_heading(doc, "ПРЕТЕНЗИЯ", size=16)
    _add_centered_heading(doc, "{тема_претензии}", size=14, bold=False)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{дата_события}").italic = True
    p.add_run(" между мной и Вами был заключен договор ")
    p.add_run("{вид_договора}").italic = True
    p.add_run(" № ")
    p.add_run("{номер_договора}").italic = True
    p.add_run(" от ")
    p.add_run("{дата_договора}").italic = True
    p.add_run(".")

    _add_section_heading(doc, "1. Суть нарушения")
    _add_placeholder_paragraph(doc, "{описание_нарушения}")

    _add_section_heading(doc, "2. Правовое основание")
    _add_placeholder_paragraph(doc, "{ссылки_на_законы_и_практику}")

    _add_section_heading(doc, "3. Требования")
    _add_placeholder_paragraph(doc, "{перечень_требований}")

    _add_section_heading(doc, "4. Срок рассмотрения")
    p = doc.add_paragraph(
        "Прошу рассмотреть настоящую претензию и удовлетворить изложенные требования в срок до "
    )
    p.add_run("{срок_ответа}").italic = True
    p.add_run(
        ". В случае неудовлетворения требований в указанный срок я буду вынужден(а) обратиться в суд."
    )

    _add_section_heading(doc, "5. Приложения")
    _add_placeholder_paragraph(doc, "{перечень_приложений}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run("{дата_составления}    ")
    p.add_run("Подпись: _______________ ").bold = True
    p.add_run("{ФИО_истец}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_isk_template(output_path: Path) -> Path:
    """Create isk statement DOCX template."""
    doc = Document()

    _add_centered_heading(doc, "В {наименование_суда}", size=14, bold=False)
    _add_placeholder_paragraph(doc, "{адрес_суда}")
    _add_placeholder_paragraph(doc, "Истец: {ФИО_истец}")
    _add_placeholder_paragraph(doc, "Адрес: {адрес_истец}")
    _add_placeholder_paragraph(doc, "Ответчик: {наименование_ответчик}")
    _add_placeholder_paragraph(doc, "Адрес: {адрес_ответчик}")
    _add_placeholder_paragraph(doc, "Цена иска: {цена_иска} руб.")
    _add_placeholder_paragraph(doc, "Госпошлина: {госпошлина} руб.")

    doc.add_paragraph()
    _add_centered_heading(doc, "ИСКОВОЕ ЗАЯВЛЕНИЕ", size=16)
    _add_centered_heading(doc, "{тема_иска}", size=14, bold=False)

    _add_section_heading(doc, "1. Обстоятельства дела")
    _add_placeholder_paragraph(doc, "{описание_обстоятельств}")

    _add_section_heading(doc, "2. Доказательства")
    _add_placeholder_paragraph(doc, "{перечень_доказательств}")

    _add_section_heading(doc, "3. Правовое обоснование")
    _add_placeholder_paragraph(doc, "{ссылки_на_законы_и_практику}")

    _add_section_heading(doc, "4. Требования к ответчику")
    _add_placeholder_paragraph(doc, "{требования_к_ответчику}")

    _add_section_heading(doc, "5. Досудебный порядок")
    p = doc.add_paragraph(
        "Досудебный порядок урегулирования спора соблюден: направлена претензия от "
    )
    p.add_run("{дата_претензии}").italic = True
    p.add_run(" (почтовое уведомление/ответ прилагается).")

    _add_section_heading(doc, "6. Приложения")
    _add_placeholder_paragraph(doc, "{перечень_приложений}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run("{дата_составления}    ")
    p.add_run("Подпись: _______________ ").bold = True
    p.add_run("{ФИО_истец}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_dogovor_template(output_path: Path) -> Path:
    """Create contract DOCX template."""
    doc = Document()

    _add_centered_heading(doc, "ДОГОВОР № {номер_договора}", size=16)
    _add_centered_heading(doc, "{вид_договора}", size=14, bold=False)

    doc.add_paragraph()
    _add_placeholder_paragraph(doc, "{место_заключения}, {дата_заключения}")

    p = doc.add_paragraph()
    p.add_run("{наименование_сторона_1}").bold = True
    p.add_run(", именуемое в дальнейшем «Заказчик», в лице ")
    p.add_run("{должность_представителя_1} {ФИО_представителя_1}").italic = True
    p.add_run(", действующего на основании ")
    p.add_run("{основание_полномочий_1}").italic = True
    p.add_run(", с одной стороны, и ")
    p.add_run("{наименование_сторона_2}").bold = True
    p.add_run(", именуемое в дальнейшем «Исполнитель», в лице ")
    p.add_run("{должность_представителя_2} {ФИО_представителя_2}").italic = True
    p.add_run(", действующего на основании ")
    p.add_run("{основание_полномочий_2}").italic = True
    p.add_run(
        ", с другой стороны, совместно именуемые «Стороны», заключили настоящий договор о нижеследующем:"
    )

    _add_section_heading(doc, "1. Предмет договора")
    _add_placeholder_paragraph(doc, "{предмет_договора}")

    _add_section_heading(doc, "2. Права и обязанности сторон")
    _add_placeholder_paragraph(doc, "{права_и_обязанности_сторон}")

    _add_section_heading(doc, "3. Стоимость работ и порядок расчетов")
    _add_placeholder_paragraph(doc, "{стоимость_и_порядок_расчетов}")

    _add_section_heading(doc, "4. Сроки выполнения")
    _add_placeholder_paragraph(doc, "{сроки_выполнения}")

    _add_section_heading(doc, "5. Ответственность сторон")
    _add_placeholder_paragraph(doc, "{ответственность_сторон}")

    _add_section_heading(doc, "6. Порядок разрешения споров")
    _add_placeholder_paragraph(doc, "{порядок_разрешения_споров}")

    _add_section_heading(doc, "7. Заключительные положения")
    _add_placeholder_paragraph(doc, "{заключительные_положения}")

    _add_section_heading(doc, "8. Реквизиты и подписи сторон")
    _add_placeholder_paragraph(doc, "{реквизиты_и_подписи}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_soglashenie_template(output_path: Path) -> Path:
    """Create agreement (soglashenie) DOCX template."""
    doc = Document()

    _add_centered_heading(doc, "СОГЛАШЕНИЕ", size=16)
    _add_centered_heading(doc, "{тема_соглашения}", size=14, bold=False)

    doc.add_paragraph()
    _add_placeholder_paragraph(doc, "{место_заключения}, {дата_заключения}")

    p = doc.add_paragraph()
    p.add_run("{наименование_сторона_1}").bold = True
    p.add_run(", именуемое в дальнейшем «Сторона 1», в лице ")
    p.add_run("{должность_представителя_1} {ФИО_представителя_1}").italic = True
    p.add_run(", и ")
    p.add_run("{наименование_сторона_2}").bold = True
    p.add_run(", именуемое в дальнейшем «Сторона 2», в лице ")
    p.add_run("{должность_представителя_2} {ФИО_представителя_2}").italic = True
    p.add_run(
        ", совместно именуемые «Стороны», заключили настоящее соглашение о нижеследующем:"
    )

    _add_section_heading(doc, "1. Предмет соглашения")
    _add_placeholder_paragraph(doc, "{предмет_соглашения}")

    _add_section_heading(doc, "2. Условия соглашения")
    _add_placeholder_paragraph(doc, "{условия_соглашения}")

    _add_section_heading(doc, "3. Срок действия")
    _add_placeholder_paragraph(doc, "{срок_действия}")

    _add_section_heading(doc, "4. Ответственность сторон")
    _add_placeholder_paragraph(doc, "{ответственность_сторон}")

    _add_section_heading(doc, "5. Заключительные положения")
    _add_placeholder_paragraph(doc, "{заключительные_положения}")

    _add_section_heading(doc, "6. Подписи сторон")
    _add_placeholder_paragraph(doc, "{подписи_сторон}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_zaklyuchenie_template(output_path: Path) -> Path:
    """Create legal opinion DOCX template."""
    doc = Document()

    _add_centered_heading(doc, "ПРАВОВОЕ ЗАКЛЮЧЕНИЕ", size=16)
    _add_centered_heading(doc, "{тема_заключения}", size=14, bold=False)

    _add_placeholder_paragraph(doc, "Дата составления: {дата_составления}")
    _add_placeholder_paragraph(doc, "Заказчик: {наименование_заказчик}")
    _add_placeholder_paragraph(doc, "Исполнитель: {ФИО_исполнитель}")

    _add_section_heading(doc, "1. Исходные данные и вопросы")
    _add_placeholder_paragraph(doc, "{исходные_данные_и_вопросы}")

    _add_section_heading(doc, "2. Применимое законодательство")
    _add_placeholder_paragraph(doc, "{ссылки_на_законы}")

    _add_section_heading(doc, "3. Судебная практика")
    _add_placeholder_paragraph(doc, "{ссылки_на_практику}")

    _add_section_heading(doc, "4. Правовой анализ")
    _add_placeholder_paragraph(doc, "{правовой_анализ}")

    _add_section_heading(doc, "5. Выводы и рекомендации")
    _add_placeholder_paragraph(doc, "{выводы_и_рекомендации}")

    _add_section_heading(doc, "6. Риски")
    _add_placeholder_paragraph(doc, "{риски}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Подготовил(а): ").bold = True
    p.add_run("{ФИО_исполнитель}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_zhaloba_template(output_path: Path) -> Path:
    """Create complaint DOCX template."""
    doc = Document()

    _add_placeholder_paragraph(doc, "{наименование_органа}")
    _add_placeholder_paragraph(doc, "{адрес_органа}")
    _add_placeholder_paragraph(doc, "От {ФИО_заявитель}")
    _add_placeholder_paragraph(doc, "Адрес: {адрес_заявитель}")

    doc.add_paragraph()
    _add_centered_heading(doc, "ЖАЛОБА", size=16)
    _add_centered_heading(doc, "{тема_жалобы}", size=14, bold=False)

    _add_section_heading(doc, "1. Суть обращения")
    _add_placeholder_paragraph(doc, "{описание_ситуации}")

    _add_section_heading(doc, "2. Нарушенные права")
    _add_placeholder_paragraph(doc, "{нарушенные_права}")

    _add_section_heading(doc, "3. Правовое основание")
    _add_placeholder_paragraph(doc, "{ссылки_на_законы}")

    _add_section_heading(doc, "4. Прошу")
    _add_placeholder_paragraph(doc, "{требования}")

    _add_section_heading(doc, "5. Приложения")
    _add_placeholder_paragraph(doc, "{перечень_приложений}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run("{дата_составления}    ")
    p.add_run("Подпись: _______________ ").bold = True
    p.add_run("{ФИО_заявитель}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def create_otvet_na_pretensiyu_template(output_path: Path) -> Path:
    """Create response to pretension DOCX template."""
    doc = Document()

    _add_placeholder_paragraph(doc, "{наименование_истец}")
    _add_placeholder_paragraph(doc, "{адрес_истец}")
    _add_placeholder_paragraph(doc, "От {наименование_ответчик}")
    _add_placeholder_paragraph(doc, "Адрес: {адрес_ответчик}")

    doc.add_paragraph()
    _add_centered_heading(doc, "ОТВЕТ НА ПРЕТЕНЗИЮ", size=16)
    _add_centered_heading(
        doc, "от {дата_претензии} № {номер_претензии}", size=14, bold=False
    )

    _add_section_heading(doc, "1. Суть полученной претензии")
    _add_placeholder_paragraph(doc, "{суть_претензии}")

    _add_section_heading(doc, "2. Позиция ответчика")
    _add_placeholder_paragraph(doc, "{позиция_ответчика}")

    _add_section_heading(doc, "3. Правовое обоснование")
    _add_placeholder_paragraph(doc, "{ссылки_на_законы_и_практику}")

    _add_section_heading(doc, "4. Решение по претензии")
    _add_placeholder_paragraph(doc, "{решение_по_претензии}")

    _add_section_heading(doc, "5. Приложения")
    _add_placeholder_paragraph(doc, "{перечень_приложений}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run("{дата_составления}    ")
    p.add_run("Подпись: _______________ ").bold = True
    p.add_run("{ФИО_представителя_ответчика}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


TEMPLATES = {
    "Претензия.docx": create_pretension_template,
    "Исковое_заявление.docx": create_isk_template,
    "Договор.docx": create_dogovor_template,
    "Соглашение.docx": create_soglashenie_template,
    "Правовое_заключение.docx": create_zaklyuchenie_template,
    "Жалоба.docx": create_zhaloba_template,
    "Ответ_на_претензию.docx": create_otvet_na_pretensiyu_template,
}


def generate_all_templates(templates_dir: Path) -> list[Path]:
    """Generate all DOCX templates and return their paths."""
    templates_dir.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    for filename, builder in TEMPLATES.items():
        output_path = templates_dir / filename
        builder(output_path)
        generated.append(output_path)
    return generated


if __name__ == "__main__":
    generated = generate_all_templates(Path("results/templates"))
    for path in generated:
        print(f"Generated: {path}")
