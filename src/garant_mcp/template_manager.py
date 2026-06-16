"""Template management for Garant lawyer agent."""

import json
import logging
from pathlib import Path
from typing import Any, Optional

from .config import BASE_DIR

logger = logging.getLogger(__name__)

TEMPLATES_DIR = BASE_DIR / "results/templates"
INDEX_FILE = TEMPLATES_DIR / "templates_index.json"


def _load_index() -> dict[str, Any]:
    """Load templates index from JSON."""
    if not INDEX_FILE.exists():
        return {"version": "1.0.0", "templates": []}
    with open(INDEX_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        return {"version": "1.0.0", "templates": []}
    return data


def find_template(request: str) -> Optional[dict[str, Any]]:
    """Find best matching template by user request keywords.

    Args:
        request: User request or case description.

    Returns:
        Template info dict or None if no match found.
    """
    index = _load_index()
    templates = index.get("templates", [])
    if not isinstance(templates, list):
        return None

    request_lower = request.lower()
    best_match: Optional[dict[str, Any]] = None
    best_score = 0

    for template in templates:
        if not isinstance(template, dict):
            continue
        score = sum(
            1
            for keyword in template.get("keywords", [])
            if isinstance(keyword, str) and keyword.lower() in request_lower
        )
        if score > best_score:
            best_score = score
            best_match = template

    if best_match is None and templates:
        first = templates[0]
        if isinstance(first, dict):
            best_match = first

    if best_match:
        logger.info(
            f"Template match for '{request[:50]}...': "
            f"{best_match['name']} (score={best_score})"
        )

    return best_match


def list_templates() -> list[dict[str, Any]]:
    """Return list of available templates."""
    index = _load_index()
    templates = index.get("templates", [])
    if not isinstance(templates, list):
        return []
    return [t for t in templates if isinstance(t, dict)]


def get_template_path(template_name: str) -> Path:
    """Return full path to template file.

    Accepts exact filename (e.g. 'Претензия.docx') or partial request.
    If partial request is given, tries to find the best matching template.
    """
    # Try exact filename first
    path = TEMPLATES_DIR / template_name
    if path.exists():
        return path

    # Try to find by keywords
    index = _load_index()
    templates = index.get("templates", [])
    if not isinstance(templates, list):
        raise FileNotFoundError(f"Template not found: {template_name}")

    request_lower = template_name.lower()
    best_match: Optional[dict[str, Any]] = None
    best_score = 0
    for item in templates:
        if not isinstance(item, dict):
            continue
        template_name_raw = item.get("name", "")
        if not isinstance(template_name_raw, str):
            continue
        name_lower = template_name_raw.lower()
        score = 0
        if name_lower == request_lower:
            score = 1000
        elif name_lower.replace(".docx", "") == request_lower.replace(".docx", ""):
            score = 900
        else:
            score = sum(
                1
                for keyword in item.get("keywords", [])
                if isinstance(keyword, str) and keyword.lower() in request_lower
            )
            if request_lower in name_lower:
                score += 50
        if score > best_score:
            best_score = score
            best_match = item

    if best_match:
        matched_name = best_match.get("name")
        if isinstance(matched_name, str):
            found_path = TEMPLATES_DIR / matched_name
            if found_path.exists():
                logger.info(
                    f"Resolved template '{template_name}' -> {matched_name} "
                    f"(score={best_score})"
                )
                return found_path

    raise FileNotFoundError(f"Template not found: {template_name}")


def resolve_output_path(
    output_path: Optional[str],
    case_path: Optional[str],
    template_name: str,
) -> Path:
    """Resolve output path for rendered document.

    Priority:
    1. Explicit output_path if given and not empty.
    2. case_path/результат/<template_name> if case_path given.
    3. results/temp/<template_name> as fallback.
    """
    if output_path:
        return Path(output_path)

    if case_path:
        case = Path(case_path)
        result_dir = case / "результат"
        result_dir.mkdir(parents=True, exist_ok=True)
        return result_dir / template_name

    fallback = BASE_DIR / "results" / "temp"
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback / template_name


def render_template(
    template_query: str,
    output_path: Optional[Path] = None,
    placeholders: Optional[dict[str, str]] = None,
    case_path: Optional[str] = None,
) -> Path:
    """Render DOCX template by replacing placeholders.

    Placeholders in template are wrapped in curly braces, e.g. {ФИО_истец}.

    Args:
        template_query: Exact template filename (e.g. 'Претензия.docx') or
            a request like 'претензия' to auto-resolve template.
        output_path: Where to save rendered document. If omitted, saves to
            case_path/результат/ or results/temp/.
        placeholders: Mapping placeholder name -> value.
        case_path: Optional case folder path. Used to determine output path.

    Returns:
        Path to rendered DOCX.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is required for rendering templates") from e

    placeholders = placeholders or {}
    template_path = get_template_path(template_query)
    template_name = template_path.name
    output = output_path or resolve_output_path(None, case_path, template_name)

    doc = Document(str(template_path))

    def replace_in_text(text: str) -> str:
        for key, value in placeholders.items():
            text = text.replace(f"{{{key}}}", str(value))
        return text

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.runs:
            full_text = "".join(run.text for run in paragraph.runs)
            new_text = replace_in_text(full_text)
            if new_text != full_text:
                first_run = paragraph.runs[0]
                first_run.text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        full_text = "".join(run.text for run in paragraph.runs)
                        new_text = replace_in_text(full_text)
                        if new_text != full_text:
                            first_run = paragraph.runs[0]
                            first_run.text = new_text
                            for run in paragraph.runs[1:]:
                                run.text = ""

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    logger.info(f"Rendered template {template_name} to {output}")
    return output


def copy_template_to_case(
    template_name: str,
    case_path: str,
    new_filename: Optional[str] = None,
) -> str:
    """Copy template DOCX to case result folder.

    Args:
        template_name: Template file name.
        case_path: Path to case folder.
        new_filename: Optional new filename.

    Returns:
        Path to copied template.
    """
    import shutil

    template_path = get_template_path(template_name)
    case = Path(case_path)
    result_dir = case / "результат"
    result_dir.mkdir(parents=True, exist_ok=True)

    dest = result_dir / (new_filename or template_path.name)
    shutil.copy2(template_path, dest)
    return str(dest)
